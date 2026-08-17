from __future__ import annotations

import io
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.oxml.parser import parse_xml
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image, ImageOps


# These colors are the colors used for the three disciplines in the source report.
DISCIPLINE_COLORS = {
    "Bouwkundig": "FBFF4A",
    "Installatietechnisch": "05BD00",
    "Organisatorisch": "002CBD",
}

BUILDING_FIELDS = [
    "complexnummer", "complexnaam", "projectadres", "plaats", "bouwjaar",
    "aantal_bouwlagen", "aantal_woningen", "grondgebonden", "portiek",
    "galerij", "corridor", "atrium", "lift",
]
INSTALLATION_FIELDS = [
    "brandmeldinstallatie", "overige_installaties", "bouwjaar_installatie", "pve",
    "onderhoud_bmi_oai", "onderhoud_blusmiddelen", "onderhoud_noodverlichting",
    "drukgeregelde_ventilatie", "zelfregelende_ventielen",
]
ORGANISATION_FIELDS = [
    "type_bezit", "woonvorm", "zorgzwaarte", "demarcatie",
    "melding_brandveilig_gebruik", "bhv", "ontruimingsplan",
    "ontruimingsplattegronden",
]
MATERIAL_FIELDS = [
    "bouwconstructie", "dakconstructie", "dakisolatie", "dakbedekking", "gevels",
    "gevelisolatie", "scheidingswanden", "vloeren", "vloerafwerking",
    "verlaagde_plafonds", "buitenkozijnen", "binnenkozijnen", "buitentrappen",
]


def _text(element) -> str:
    return "".join(element.xpath(".//w:t/text()"))


def _find_paragraph(doc: Document, exact: str):
    for paragraph in doc.paragraphs:
        if " ".join(paragraph.text.split()) == exact:
            return paragraph
    raise ValueError(f"Template paragraph not found: {exact}")


def _replace_text_nodes(element, replacements: dict[str, str]) -> None:
    for node in element.xpath(".//w:t"):
        value = node.text or ""
        for old, new in replacements.items():
            value = value.replace(old, str(new or ""))
        node.text = value


def _set_paragraph_text(paragraph, value: str) -> None:
    nodes = paragraph._p.xpath(".//w:t")
    if nodes:
        nodes[0].text = str(value or "")
        for node in nodes[1:]:
            node.text = ""
    else:
        paragraph.add_run(str(value or ""))


def _set_label_value_paragraph(paragraph, label: str, value: str) -> None:
    p = paragraph._p
    source_run = next((r for r in p.xpath(".//w:r") if r.xpath(".//w:t")), None)
    rpr = deepcopy(source_run.find(qn("w:rPr"))) if source_run is not None and source_run.find(qn("w:rPr")) is not None else None
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    run = OxmlElement("w:r")
    if rpr is not None:
        run.append(rpr)
    label_node = OxmlElement("w:t")
    label_node.text = label
    run.append(label_node)
    run.append(OxmlElement("w:tab"))
    value_node = OxmlElement("w:t")
    value_node.text = str(value or "")
    run.append(value_node)
    p.append(run)


def _new_paragraph_from(source, text: str):
    new_p = deepcopy(source._p)
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    run = OxmlElement("w:r")
    source_runs = source._p.xpath(".//w:r[w:t]")
    if source_runs:
        rpr = source_runs[0].find(qn("w:rPr"))
        if rpr is not None:
            run.append(deepcopy(rpr))
    parts = str(text or "").split("\n")
    for index, part in enumerate(parts):
        if index:
            run.append(OxmlElement("w:br"))
        text_node = OxmlElement("w:t")
        if part.startswith(" ") or part.endswith(" "):
            text_node.set(qn("xml:space"), "preserve")
        text_node.text = part
        run.append(text_node)
    new_p.append(run)
    return new_p


def _set_cell_lines(cell, lines: list[str], source_paragraph=None) -> None:
    tc = cell._tc
    paragraphs = cell.paragraphs
    if not paragraphs:
        seed = OxmlElement("w:p")
        tc.append(seed)
        paragraphs = [Paragraph(seed, cell)]
    source = source_paragraph or next((p for p in paragraphs if p.text.strip()), paragraphs[0])
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)
    clean_lines = [str(line or "") for line in lines] or [""]
    for line in clean_lines:
        tc.append(_new_paragraph_from(source, line))


def _set_cell_value(cell, value: str) -> None:
    lines = str(value or "").replace("\r", "").split("\n")
    _set_cell_lines(cell, lines)


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _shade_row(row, fill: str) -> None:
    """Apply shading to all cells in a row."""
    for cell in row.cells:
        _shade(cell, fill)


def _remove_between(body, start_element, end_element, include_start=False) -> None:
    children = list(body)
    start = children.index(start_element) + (0 if include_start else 1)
    end = children.index(end_element)
    for child in children[start:end]:
        body.remove(child)


def _page_break_paragraph() -> OxmlElement:
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run.append(br)
    paragraph.append(run)
    return paragraph


def _spacer_paragraph(style_id: str = "Geenafstand1") -> OxmlElement:
    paragraph = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    p_style = OxmlElement("w:pStyle")
    p_style.set(qn("w:val"), style_id)
    p_pr.append(p_style)
    paragraph.append(p_pr)
    return paragraph


def _fit_image_bytes(path: str | Path, width_px: int, height_px: int, crop=False) -> io.BytesIO:
    with Image.open(path) as image:
        image = ImageOps.exif_transpose(image).convert("RGB")
        if crop:
            image = ImageOps.fit(image, (width_px, height_px), method=Image.Resampling.LANCZOS)
        else:
            image.thumbnail((width_px, height_px), Image.Resampling.LANCZOS)
        stream = io.BytesIO()
        image.save(stream, "JPEG", quality=92, optimize=True)
        stream.seek(0)
        return stream


def _clear_cell_keep_properties(cell) -> None:
    tc = cell._tc
    p_pr = deepcopy(cell.paragraphs[0]._p.pPr) if cell.paragraphs and cell.paragraphs[0]._p.pPr is not None else None
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)
    p = OxmlElement("w:p")
    if p_pr is not None:
        p.append(p_pr)
    tc.append(p)


def _add_cell_photo(cell, path: str | None, width_cm=7.15, max_height_cm=5.45) -> None:
    _clear_cell_keep_properties(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if not path or not Path(path).exists():
        return
    with Image.open(path) as image:
        image = ImageOps.exif_transpose(image)
        ratio = image.height / max(image.width, 1)
    height_cm = min(max_height_cm, width_cm * ratio)
    width_used = min(width_cm, height_cm / max(ratio, 0.001))
    stream = _fit_image_bytes(path, 1500, 1100, crop=False)
    p.add_run().add_picture(stream, width=Cm(width_used), height=Cm(height_cm))


def _inline_to_anchor(inline, x_cm: float, y_cm: float) -> None:
    anchor = deepcopy(inline)
    anchor.tag = qn("wp:anchor")
    for key, value in {
        "distT": "0", "distB": "0", "distL": "0", "distR": "0",
        "simplePos": "0", "relativeHeight": "251659264", "behindDoc": "0",
        "locked": "0", "layoutInCell": "1", "allowOverlap": "1",
    }.items():
        anchor.set(key, value)
    for child in list(anchor):
        anchor.remove(child)
    simple = OxmlElement("wp:simplePos")
    simple.set("x", "0")
    simple.set("y", "0")
    pos_h = OxmlElement("wp:positionH")
    pos_h.set("relativeFrom", "page")
    off_h = OxmlElement("wp:posOffset")
    off_h.text = str(int(Cm(x_cm)))
    pos_h.append(off_h)
    pos_v = OxmlElement("wp:positionV")
    pos_v.set("relativeFrom", "page")
    off_v = OxmlElement("wp:posOffset")
    off_v.text = str(int(Cm(y_cm)))
    pos_v.append(off_v)
    old_children = list(inline)
    anchor.extend([simple, pos_h, pos_v])
    for child in old_children:
        if child.tag == qn("wp:extent"):
            anchor.append(deepcopy(child))
        elif child.tag == qn("wp:effectExtent"):
            anchor.append(deepcopy(child))
            anchor.append(OxmlElement("wp:wrapNone"))
        else:
            anchor.append(deepcopy(child))
    inline.getparent().replace(inline, anchor)


def _set_cover_photo(doc: Document, photo_path: str | None) -> None:
    paragraph = doc.paragraphs[16]
    for drawing in list(paragraph._p.xpath(".//w:drawing|.//w:pict")):
        parent = drawing.getparent()
        if parent is not None:
            parent.remove(drawing)
    for node in paragraph._p.xpath(".//w:t"):
        node.text = ""
    if not photo_path or not Path(photo_path).exists():
        return
    stream = _fit_image_bytes(photo_path, 1600, 1540, crop=True)
    inline = paragraph.add_run().add_picture(stream, width=Cm(13.9), height=Cm(13.4))._inline
    _inline_to_anchor(inline, 3.40, 7.75)


def _set_content_control_value(control, value: str) -> None:
    """Replace only the editable text nodes inside an existing template control."""
    nodes = list(control.iter(qn("w:t")))
    if not nodes:
        return
    nodes[0].text = str(value or "")
    for node in nodes[1:]:
        node.text = ""


def _fill_data_table(table, keys: list[str], project: dict) -> None:
    """Keep template labels intact and write exclusively into the value slot.

    Several rows in 3.2, 3.3 and 3.4 are visually two-column rows but contain
    a single physical cell plus a Word content control. Rebuilding those rows
    erased their fixed labels. This routine edits the existing content control
    when present and otherwise edits only the physical right-hand cell.
    """
    if len(table.rows) != len(keys):
        raise ValueError(f"Template table has {len(table.rows)} rows; expected {len(keys)}")
    for row, key in zip(table.rows, keys):
        value = project.get(key, "")
        controls = row._tr.xpath(".//w:sdt")
        if controls:
            _set_content_control_value(controls[-1], value)
        elif len(row._tr.tc_lst) >= 2:
            _set_cell_value(row.cells[-1], value)
        else:
            raise ValueError(f"No editable value slot found for template field {key}")
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def _value_lines(value: str) -> list[str]:
    value = str(value or "").replace("\r", "").strip()
    return value.split("\n") if value else [""]


def _set_horizontal_border(cell, edge: str, value: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    border = borders.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        borders.append(border)
    border.set(qn("w:val"), value)
    border.set(qn("w:sz"), "4")
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), "auto")


def _replace_detail_row_with_pairs(table: Table, pairs: list[tuple[str, str]]) -> None:
    """Create coupled label/value rows so multi-line values cannot drift vertically."""
    source = table.rows[4]._tr
    parent = source.getparent()
    position = parent.index(source)
    parent.remove(source)
    for offset in range(len(pairs)):
        parent.insert(position + offset, deepcopy(source))

    detail_rows = table.rows[4:4 + len(pairs)]
    for index, (row, pair) in enumerate(zip(detail_rows, pairs)):
        label, value = pair
        _set_cell_value(row.cells[0], label)
        _set_cell_value(row.cells[1], value)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            _set_horizontal_border(cell, "top", "single" if index == 0 else "nil")
            _set_horizontal_border(cell, "bottom", "single" if index == len(pairs) - 1 else "nil")


def _fill_finding_table(table: Table, finding: dict, other=False) -> None:
    code = f"{finding.get('code_group', 'G')}.{int(finding.get('code_number') or 0):02d}"
    if other:
        _set_cell_value(table.rows[0].cells[0], "Foto tijdens inspectie")
        _shade_row(table.rows[0], "A6A6A6")
        table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
        _set_cell_value(table.rows[1].cells[0], "Overzichtsfoto")
        _set_cell_value(table.rows[1].cells[1], "Detailfoto")
    else:
        _set_cell_value(table.rows[0].cells[0], f"Gebrek {code}")
        _shade_row(table.rows[0], DISCIPLINE_COLORS.get(finding.get("discipline"), "FBFF4A"))
    _add_cell_photo(table.rows[2].cells[0], finding.get("photo_before"))
    _add_cell_photo(table.rows[2].cells[1], finding.get("photo_after"))
    _set_cell_value(table.rows[3].cells[0], "Algemene informatie")
    _shade_row(table.rows[3], "D9D9D9")

    if other:
        labels = ["Tekeningnummer:", "Bouwlaag:", "Ruimtenummer:", "Eis:", "Bevinding:", "Conclusie:"]
        values = [
            finding.get("tekeningnummer"), finding.get("bouwlaag"), finding.get("ruimte"),
            finding.get("eis"), finding.get("gebrek"), finding.get("maatregel"),
        ]
    else:
        labels = ["Tekeningnummer:", "Bouwlaag:", "Ruimtenummer:", "Eis:", "Gebrek:", "Aantal:", "Afmeting:", "Maatregel:", "Opmerking:"]
        values = [
            finding.get("tekeningnummer"), finding.get("bouwlaag"), finding.get("ruimte"),
            finding.get("eis"), finding.get("gebrek"), finding.get("aantal"),
            finding.get("afmeting"), finding.get("maatregel"), finding.get("opmerking"),
        ]
    _replace_detail_row_with_pairs(
        table,
        [(label, "\n".join(_value_lines(value))) for label, value in zip(labels, values)],
    )


def _insert_findings(doc: Document, findings: list[dict]) -> None:
    body = doc._element.body
    measures_heading = _find_paragraph(doc, "Maatregelen")
    measures_last = _find_paragraph(doc, "De benoemde hoeveelheden moeten per gebrek in het werk gecontroleerd worden.")
    other_heading = _find_paragraph(doc, "Overige bevindingen")
    equality_heading = _find_paragraph(doc, "Gelijkwaardigheidsoplossing")
    measure_template = deepcopy(doc.tables[7]._tbl)
    completion_template = deepcopy(doc.tables[8]._tbl)
    other_template = deepcopy(doc.tables[31]._tbl)

    _remove_between(body, measures_last._p, other_heading._p)
    _remove_between(body, other_heading._p, equality_heading._p)

    insert_at = list(body).index(other_heading._p)
    measures = [f for f in findings if f.get("finding_type") == "Maatregel"]
    for finding in measures:
        body.insert(insert_at, _page_break_paragraph())
        insert_at += 1
        table_el = deepcopy(measure_template)
        body.insert(insert_at, table_el)
        _fill_finding_table(Table(table_el, doc), finding, other=False)
        insert_at += 1
        body.insert(insert_at, _spacer_paragraph())
        insert_at += 1
        completion_el = deepcopy(completion_template)
        body.insert(insert_at, completion_el)
        insert_at += 1

    insert_at = list(body).index(equality_heading._p)
    others = [f for f in findings if f.get("finding_type") == "Overige bevinding"]
    for finding in others:
        body.insert(insert_at, _spacer_paragraph("Normal"))
        insert_at += 1
        table_el = deepcopy(other_template)
        body.insert(insert_at, table_el)
        _fill_finding_table(Table(table_el, doc), finding, other=True)
        insert_at += 1


def _fill_situation_table(doc: Document, project: dict) -> None:
    keys = ["photo_voorgevel", "photo_kopgevel", "photo_achtergevel", "photo_luchtfoto"]
    table = doc.tables[6]
    for index, key in enumerate(keys):
        cell = table.cell(index // 2, index % 2)
        caption = ["Voorgevel", "Kopgevel", "Achtergevel", "Luchtfoto (bron: google.maps)"][index]
        _clear_cell_keep_properties(cell)
        _add_cell_photo(cell, project.get(key), width_cm=7.2, max_height_cm=4.55)
        paragraph = cell.add_paragraph(caption)
        paragraph.style = "Geen afstand1"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _fill_front_matter(doc: Document, project: dict) -> None:
    _set_paragraph_text(doc.paragraphs[4], project.get("complexnaam") or project.get("title") or "")
    _set_label_value_paragraph(doc.paragraphs[33], "Kenmerk:", project.get("kenmerk", ""))
    _set_label_value_paragraph(doc.paragraphs[34], "Datum:", project.get("report_date", ""))
    _set_label_value_paragraph(doc.paragraphs[35], "Versie:", project.get("version", ""))

    table = doc.tables[0]
    values = [
        [project.get("projectadres", ""), " ".join(x for x in [project.get("postcode", ""), project.get("plaats", "")] if x)],
        [project.get("opdrachtgever", ""), project.get("opdrachtgever_adres", "")],
        [project.get("opdrachtgever_contact", "")],
        ["TriaCon BV", "Postbus 40064", "8004 DB Zwolle", "Tel: 038-3338070"],
        _value_lines(project.get("inspecteur", "")),
        _value_lines(project.get("gecontroleerd", "")),
        [project.get("triacon_contact", ""), f"E-mail: {project.get('triacon_email', '')}" if project.get("triacon_email") else ""],
    ]
    for row, lines in zip(table.rows, values):
        _set_cell_lines(row.cells[1], [x for x in lines if x != ""] or [""])
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    version_table = doc.tables[1]
    _set_cell_value(version_table.rows[1].cells[0], f"Versie {project.get('version', '')}")
    _set_cell_value(version_table.rows[1].cells[1], project.get("report_date", ""))
    _set_cell_value(version_table.rows[1].cells[2], "")


def _fill_fixed_and_variable_text(doc: Document, project: dict) -> None:
    # The paragraphs below remain the actual template paragraphs. Only their documented
    # placeholders are replaced, so the standard wording and run formatting stay intact.
    _replace_text_nodes(doc._element, {
        "[Naam opdrachtgever]": project.get("opdrachtgever", ""),
        "[Naam complex]": project.get("complexnaam", ""),
        "[Complex naam]": project.get("complexnaam", ""),
        "[Datum Complexnummer]": " ".join(x for x in [project.get("report_date", ""), project.get("complexnummer", "")] if x),
    })
    appendix_prefix = " ".join(x for x in [project.get("report_date", ""), project.get("complexnummer", "")] if x)
    for paragraph in doc.paragraphs:
        if "[Datum Complexnummer]" in paragraph.text:
            _set_paragraph_text(paragraph, paragraph.text.replace("[Datum Complexnummer]", appendix_prefix))
    complex_code = project.get("complexnummer", "") or "…"
    code_paragraph = _find_paragraph(doc, "Complexnummer (…);")
    _replace_text_nodes(code_paragraph._p, {"…": complex_code})

    _set_paragraph_text(_find_paragraph(doc, "[Invullen]"), project.get("algemene_omschrijving") or "[Nog invullen]")
    visited = _find_paragraph(doc, "[Invullen welke woningen zijn bezocht]")
    _set_paragraph_text(visited, project.get("bezochte_woningen") or "[Nog invullen]")
    remaining = [p for p in doc.paragraphs if " ".join(p.text.split()) == "[Invullen]"]
    if remaining:
        _set_paragraph_text(remaining[-1], project.get("beperkingen") or "[Nog invullen]")

    # Replace tekeningen_ontvangen value in Uitgangspunten section
    tekeningen_waarde = project.get("tekeningen_ontvangen", "geen")
    if tekeningen_waarde and tekeningen_waarde != "Kies een item.":
        # Find and replace in the Uitgangspunten paragraph (paragraph 106)
        try:
            uitgangspunten_para = doc.paragraphs[106]
            if "Voorafgaand aan de inspectie zijn geen tekeningen ontvangen" in uitgangspunten_para.text:
                new_text = uitgangspunten_para.text.replace("zijn geen tekeningen", f"zijn {tekeningen_waarde} tekeningen")
                _set_paragraph_text(uitgangspunten_para, new_text)
        except (IndexError, AttributeError):
            pass  # If paragraph not found, continue

    summary_heading = _find_paragraph(doc, "Samenvatting")
    summary_source = doc.paragraphs[83]
    summary_heading._p.addnext(_new_paragraph_from(summary_source, project.get("samenvatting") or "[Nog invullen]"))
    conclusion_heading = _find_paragraph(doc, "Conclusie")
    next_el = conclusion_heading._p.getnext()
    if next_el is not None and next_el.tag == qn("w:p"):
        from docx.text.paragraph import Paragraph
        _set_paragraph_text(Paragraph(next_el, doc), project.get("conclusie") or "[Nog invullen]")
    equality_default = _find_paragraph(doc, "In dit complex zijn geen gelijkwaardigheidsoplossingen van toepassing.")
    if project.get("gelijkwaardigheid"):
        _set_paragraph_text(equality_default, project["gelijkwaardigheid"])


def _set_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def get_standard_texts(template_path: Path) -> dict[str, str]:
    doc = Document(template_path)
    return {
        "inleiding": "\n\n".join(doc.paragraphs[i].text for i in [79, 80, 81, 83]),
        "doelstelling": doc.paragraphs[103].text,
        "uitgangspunten": doc.paragraphs[106].text,
        "maatregelen": "\n\n".join(doc.paragraphs[i].text for i in [134, 136, 137, 138, 139, 140, 142, 143, 144, 145, 147]),
    }


def build_report(template_path: Path, project: dict, findings: list[dict]) -> bytes:
    doc = Document(template_path)
    _fill_front_matter(doc, project)
    _set_cover_photo(doc, project.get("photo_voorgevel"))
    _fill_fixed_and_variable_text(doc, project)
    _fill_data_table(doc.tables[2], BUILDING_FIELDS, project)
    _fill_data_table(doc.tables[3], INSTALLATION_FIELDS, project)
    _fill_data_table(doc.tables[4], ORGANISATION_FIELDS, project)
    _fill_data_table(doc.tables[5], MATERIAL_FIELDS, project)
    _fill_situation_table(doc, project)
    _insert_findings(doc, findings)
    _set_update_fields(doc)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
